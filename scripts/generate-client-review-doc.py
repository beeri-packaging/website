#!/usr/bin/env python3
"""Generate a Hebrew client content-review DOCX from the current site sources."""

from __future__ import annotations

import json
import re
import subprocess
import textwrap
import xml.etree.ElementTree as ET
from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "client-review"
OUT_FILE = OUT_DIR / "beeri-arizot-content-review-he.docx"
PYTHON_DOCX_SOURCES = [
    ROOT / "docs" / "strategy" / "beeri-packaging-seo-content-strategy.docx",
    ROOT / "docs" / "briefs" / "beeri-seo-blog-brief.docx",
]
PPTX_SOURCE = ROOT / "docs" / "presentations" / "beeri-arizot-seo-findings.pptx"

INK = RGBColor(0x1B, 0x1C, 0x1A)
CLAY = RGBColor(0x4D, 0x46, 0x32)
PURPLE = RGBColor(0x6F, 0x50, 0xA6)
MAGENTA_DEEP = RGBColor(0xA3, 0x2F, 0x7F)
CYAN_DEEP = RGBColor(0x00, 0x20, 0x20)
MUTED = RGBColor(0x5F, 0x5D, 0x5C)

FILL_BONE = "FBF9F6"
FILL_SAND = "F5F3F0"
FILL_GOLD = "FFE177"
FILL_YELLOW = "FFD400"
FILL_CYAN = "CCFFFF"
FILL_RULE = "D0C6AB"
FILL_INK = "1B1C1A"


JS_EXPORTER = r"""
const ts = require('typescript');
const fs = require('fs');
const Module = require('module');
const path = require('path');
const cwd = process.cwd();

require.extensions['.ts'] = function(module, filename) {
  const source = fs.readFileSync(filename, 'utf8');
  const out = ts.transpileModule(source, {
    compilerOptions: {
      module: ts.ModuleKind.CommonJS,
      target: ts.ScriptTarget.ES2020,
      esModuleInterop: true,
      jsx: ts.JsxEmit.ReactJSX,
    }
  }).outputText;
  module._compile(out, filename);
};

const origResolve = Module._resolveFilename;
Module._resolveFilename = function(request, parent, isMain, options) {
  if (request.startsWith('@/')) request = path.join(cwd, request.slice(2));
  return origResolve.call(this, request, parent, isMain, options);
};

const home = require('./app/content/home.ts');
const about = require('./app/content/about.ts');
const catalog = require('./app/content/catalog.ts');
const finishing = require('./app/content/finishing.ts');
const careers = require('./app/content/careers.ts');
const blog = require('./app/content/blog.ts');
const site = require('./app/content/site.ts');
const contact = require('./app/content/contact.ts');
const job = require('./app/content/jobApplication.ts');
const legal = require('./app/content/legal.ts');
const placeholder = require('./app/content/placeholder.ts');
const system = require('./app/content/system.ts');
const company = require('./app/content/company.ts');
const messages = require('./messages/he.json');

console.log(JSON.stringify({
  company: company.COMPANY,
  metadata: messages.Metadata,
  chrome: site.chromeContent.he,
  navLinks: site.navLinks,
  home: {
    copy: home.homeCopy.he,
    capabilities: home.capabilities.map((item) => ({n: item.n, ...item.he})),
    faqItems: home.faqItems.map((item) => ({n: item.n, ...item.he})),
    journeyPanels: home.journeyPanels.map((item) => ({
      key: item.key,
      tag: item.he.tag,
      title: item.he.title,
      body: item.he.body,
      link: item.he.link
    }))
  },
  about: about.aboutCopy.he,
  catalog: catalog.catalogCopy.he,
  finishing: finishing.finishingCopy.he,
  careers: careers.careersCopy.he,
  blog: {
    index: blog.blogIndexCopy.he,
    categoryLabels: Object.fromEntries(Object.entries(blog.categoryLabels).map(([key, val]) => [key, val.he])),
    posts: blog.blogPosts.map((post) => ({
      slug: post.slug,
      date: post.date,
      read: post.read.he,
      author: post.author ? post.author.he : null,
      credit: post.credit ? post.credit.he : null,
      ...post.he
    }))
  },
  contact: contact.contactCopy.he,
  jobApplication: job.jobApplicationCopy.he,
  legal: {
    privacy: legal.privacyDoc.he,
    terms: legal.termsDoc.he,
  },
  placeholder: Object.fromEntries(Object.entries(placeholder.placeholderContent).map(([key, val]) => [key, val.he])),
  system: {
    notFound: system.notFoundCopy.he,
    error: system.errorCopy.he,
  }
}, null, 2));
"""


DOC_SUMMARIES = [
    {
        "file": "docs/superpowers/specs/2026-06-04-website-foundation-design.md",
        "title": "תכנון יסודות האתר",
        "summary": [
            "האתר עובר לבסיס דו-לשוני, נגיש ומהיר, עם עברית כשפת ברירת מחדל.",
            "ההחלטות המרכזיות: Next.js App Router, רכיבי שרת כברירת מחדל, Sanity כמערכת ניהול תוכן, next-intl לנתיבי /he ו-/en, ו-Tailwind לפי טוקני המותג.",
            "מבחינת לקוח, המשמעות היא אתר שקל לנהל, עם תוכן עברי ואנגלי באותו מבנה, SEO מסודר, ביצועים גבוהים ונגישות WCAG 2.1 AA.",
        ],
    },
    {
        "file": "docs/superpowers/specs/2026-06-04-cms-full-buildout-design.md",
        "title": "בניית CMS מלאה",
        "summary": [
            "המסמך מגדיר איך להעביר תוכן לאתר מנוהל ב-Sanity בלי לשבור את מבנה העמודים הקיים.",
            "התוכן אמור להישמר במסמכים דו-לשוניים: בית, קריירה, השבחות, קטלוג, בלוג, ניווט, הודעות מערכת ותמונות.",
            "ה-Studio מאורגן כך שללקוח יהיה אזור ברור לעריכת עמודים, ניווט, פוסטים, משרות ופרטי אתר.",
        ],
    },
    {
        "file": "docs/superpowers/specs/2026-06-05-contact-inquiry-design.md",
        "title": "טופס יצירת קשר",
        "summary": [
            "הטופס נועד להפוך פניות לפרויקט לאימייל מסודר עם שדות חובה מינימליים.",
            "שדות חובה: שם מלא, טלפון וסיבת פנייה. אימייל, חברה ופרטים נוספים הם שדות רשות.",
            "הטופס כולל הודעות שגיאה בעברית, מצב שליחה, הודעת הצלחה והגנת honeypot נגד ספאם.",
        ],
    },
    {
        "file": "docs/superpowers/specs/2026-06-05-parent-company-mention-design.md",
        "title": "אזכור קבוצת דפוס בארי",
        "summary": [
            "המטרה היא להציג את בארי אריזות כחלק מקבוצת דפוס בארי בלי להפוך את עמוד הבית לעמוד היסטורי.",
            "המסר צריך להיות תדמיתי-מקצועי: מורשת דפוס, אמינות, יכולת ייצור ורקע קבוצתי.",
            "העובדות שצריך לשמור עליהן: בארי אריזות היא חלק מקבוצת דפוס בארי; הקבוצה פועלת מקיבוץ בארי; השורשים כוללים את גרפיקה בצלאל ודפוס חרט.",
        ],
    },
    {
        "file": "docs/superpowers/specs/2026-06-06-about-page-design.md",
        "title": "עמוד אודות",
        "summary": [
            "עמוד האודות מציג את בארי אריזות דרך מורשת הדפוס, המפעל ביבנה, יכולות הייצור, לקוחות ותקני איכות.",
            "מבנה העמוד כולל: הירו, קבוצת דפוס בארי, ציר זמן, נתוני מפעל, יכולות, לקוחות, תקנים וקריאה לפעולה.",
            "הנתונים המרכזיים לאישור: רישום חברה 1964, ח.פ. 520026113, כ-140 עובדים, 7,900 מ\"ר ייצור ומשרדים, 3,000 מ\"ר מרלו\"ג, תקני ISO 9001 ו-FSSC 22000.",
        ],
    },
    {
        "file": "docs/superpowers/specs/2026-06-08-scroll-entrance-animations-design.md",
        "title": "אנימציות כניסה בגלילה",
        "summary": [
            "המסמך מתאר אנימציות עדינות שמופעלות כשאלמנטים נכנסים למסך.",
            "האנימציות הן שיפור חוויית משתמש בלבד: אין שינוי בתוכן, במבנה או ב-SEO.",
            "יש תמיכה בהעדפת המשתמש להפחתת תנועה, כדי לשמור על נגישות.",
        ],
    },
    {
        "file": "docs/superpowers/plans/2026-06-04-cms-full-buildout.md",
        "title": "תוכנית יישום CMS",
        "summary": [
            "תוכנית עבודה מפורטת להעברת תוכן האתר ל-Sanity בשלבים.",
            "השלבים כוללים סכמות תוכן, שכבת נתונים, מיגרציית עמודים, העלאת תמונות וזריעת תוכן עברי ואנגלי.",
            "מבחינת אישור לקוח, המסמך מחזק שהתוכן מחולק לפי עמודים וסוגי מידע, ולא מפוזר בקוד.",
        ],
    },
    {
        "file": "docs/superpowers/plans/2026-06-05-contact-inquiry.md",
        "title": "תוכנית יישום טופס פנייה",
        "summary": [
            "תוכנית לבניית ולבדיקת טופס יצירת הקשר, כולל ולידציה, תוכן עברי ואנגלי, שליחה במייל ובדיקות.",
            "הטקסטים שמעניינים את הלקוח מופיעים בפרק טופס יצירת קשר במסמך זה.",
        ],
    },
    {
        "file": "docs/superpowers/plans/2026-06-06-about-page.md",
        "title": "תוכנית יישום עמוד אודות",
        "summary": [
            "תוכנית לבניית עמוד האודות בפועל: מודול תוכן, עמוד שרת, רכיבי עיצוב, ציר זמן, לוגואים, ניווט ובדיקות.",
            "הטקסטים שמעניינים את הלקוח מופיעים בפרק עמוד אודות במסמך זה; התוכנית עצמה היא מסמך ביצוע פנימי.",
        ],
    },
]


TERMS = [
    ("אריזות קרטון ממותגות", "אריזות קרטון שמתוכננות ומודפסות לפי מותג, מוצר ויעד שימוש, ולא מוצר מדף גנרי."),
    ("התאמה אישית", "התאמת מידות, חומר, מבנה, הדפסה והשבחות לצורכי מוצר מסוים."),
    ("תכנון מבני", "תכנון הצורה הפיזית של האריזה: קיפול, סגירה, חיזוקים, חלונות ותמיכה במוצר."),
    ("דייליין", "קובץ פריסה טכני של האריזה עם קווי חיתוך, קיפול והדבקה. זה הבסיס לעיצוב ולייצור."),
    ("שטנץ", "חיתוך צורני של הקרטון לפי מבנה מוגדר, בעזרת מבלט או תהליך דיגיטלי."),
    ("מבלט", "כלי חיתוך ייעודי שמייצר את צורת האריזה, קווי הקיפול והפתחים."),
    ("דפוס דיגיטלי", "שיטת דפוס גמישה לסדרות קצרות, פיילוטים, גרסאות משתנות ולוחות זמנים קצרים."),
    ("דפוס אופסט", "שיטת דפוס איכותית ומשתלמת בכמויות בינוניות וגדולות, עם עקביות צבע גבוהה."),
    ("השבחות דפוס", "טכניקות שמוסיפות ערך חזותי או תחושתי לאריזה: פויל, לכה, הבלטה, דיבוס ולמינציה."),
    ("פויל", "שכבה מטאלית או צבעונית שמבליטה לוגו, סימון או פרט עיצובי."),
    ("לכה סלקטיבית", "לכה שמיושמת רק באזורים נבחרים כדי להבליט פרט מסוים."),
    ("הבלטה / דיבוס", "יצירת עומק מוחשי כלפי מעלה או מטה על גבי הקרטון."),
    ("חוצץ / אינסרט", "רכיב פנימי שמייצב את המוצר בתוך האריזה ומשפר הצגה והגנה."),
    ("מרלו\"ג", "מרכז לוגיסטי לאחסון והפצה."),
    ("ISO 9001", "תקן ניהול איכות שמגדיר תהליכי עבודה עקביים ובקרת איכות."),
    ("FSSC 22000", "תקן בטיחות מזון הרלוונטי במיוחד לאריזות מזון, קוסמטיקה ופארמה."),
    ("חומר בר-מיחזור", "קרטון או חומר גלם שניתן למחזר, בהתאם לסוג החומר והציפוי."),
]

LABELS_HE = {
    "Eyebrow": "תג עליון",
    "Lead": "פתיח",
    "CTA": "קריאה לפעולה",
    "CTA ראשי": "קריאה לפעולה ראשית",
    "CTA משני": "קריאה לפעולה משנית",
    "CTA השבחות": "קריאה לפעולה - השבחות",
    "CTA קטלוג": "קריאה לפעולה - קטלוג",
    "CTA מסכם": "קריאה לפעולה מסכמת",
    "Slug": "מזהה כתובת",
    "Title": "כותרת מטא",
    "Description": "תיאור מטא",
    "Hover": "ריחוף",
    "שכבת Hover": "שכבת ריחוף",
    "נתוני Hover": "נתוני ריחוף",
    "Fallback": "גיבוי",
    "Placeholder": "טקסט לדוגמה בשדה",
    "Preview": "תצוגה מקדימה",
    "Kicker": "שורת הדגשה",
}


def label_he(label: str) -> str:
    label = str(label)
    if label in LABELS_HE:
        return LABELS_HE[label]
    for source, target in LABELS_HE.items():
        label = label.replace(source, target)
    return label


def export_site_data() -> dict:
    result = subprocess.run(
        ["node", "-e", JS_EXPORTER],
        cwd=ROOT,
        check=True,
        text=True,
        capture_output=True,
    )
    return json.loads(result.stdout)


def clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = text.replace("—", "–")
    text = text.replace("…", "...")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    rtl(p)
    run = p.add_run(clean_text(str(text)))
    run.bold = bold
    set_run_font(run, bold=bold, color=INK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_border_bottom(paragraph, color: str = FILL_RULE, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = FILL_RULE) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("בארי אריזות · מסמך סקירת תוכן · עמוד ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(p)


def rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_page_direction(section) -> None:
    sect_pr = section._sectPr
    if sect_pr.find(qn("w:bidi")) is None:
        sect_pr.append(OxmlElement("w:bidi"))


def add_paragraph(doc: Document, text: str = "", *, style: str | None = None, bold: bool = False, size: float | None = None) -> None:
    if not clean_text(text):
        return
    p = doc.add_paragraph(style=style)
    rtl(p)
    run = p.add_run(clean_text(text))
    set_run_font(run, size=size, bold=bold, color=CLAY)


def add_heading(doc: Document, text: str, level: int = 1, *, page_break: bool = False) -> None:
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(style=f"Heading {level}")
    rtl(p)
    keep_with_next(p)
    run = p.add_run(clean_text(text))
    sizes = {1: 21, 2: 15, 3: 12.5}
    colors = {1: INK, 2: PURPLE, 3: MAGENTA_DEEP}
    set_run_font(run, size=sizes.get(level, 12), bold=True, color=colors.get(level, INK))
    if level == 1:
        paragraph_border_bottom(p, color="00FFFF", size="18")
        p.paragraph_format.space_after = Pt(12)


def add_label_value(doc: Document, label: str, value: str | int | None) -> None:
    if value in (None, ""):
        return
    p = doc.add_paragraph()
    rtl(p)
    r1 = p.add_run(f"{label_he(label)}: ")
    set_run_font(r1, bold=True, color=INK)
    r2 = p.add_run(clean_text(str(value)))
    set_run_font(r2, color=CLAY)


def add_bullets(doc: Document, items) -> None:
    for item in items:
        text = clean_text(str(item))
        if not text:
            continue
        p = doc.add_paragraph(style="List Bullet")
        rtl(p)
        run = p.add_run(text)
        set_run_font(run, color=CLAY)


def add_table(doc: Document, headers, rows) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], FILL_SAND)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    for row_index, row in enumerate(table.rows[1:], start=1):
        if row_index % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "FEFDFB")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, fill: str = FILL_SAND) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=fill)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    rtl(p)
    r1 = p.add_run(clean_text(title))
    set_run_font(r1, size=11, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(3)
    p2 = cell.add_paragraph()
    rtl(p2)
    r2 = p2.add_run(clean_text(body))
    set_run_font(r2, size=10.5, color=CLAY)
    doc.add_paragraph()


def join_title(value) -> str:
    if isinstance(value, (list, tuple)):
        return " ".join(str(x) for x in value)
    return str(value)


def add_copy_block(doc: Document, title: str, fields: list[tuple[str, object]]) -> None:
    add_heading(doc, title, 3)
    for label, value in fields:
        if value is None or value == "":
            continue
        if isinstance(value, (list, tuple)):
            value = " / ".join(str(x) for x in value)
        add_label_value(doc, label, value)


def read_markdown_text(path: Path) -> list[str]:
    if not path.exists():
        return []
    lines = path.read_text(encoding="utf-8").splitlines()
    output = []
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if re.fullmatch(r"-{3,}", line):
            continue
        if re.fullmatch(r"\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?", line):
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        if "|" in line:
            cells = [clean_text(cell) for cell in line.strip("|").split("|")]
            cells = [cell for cell in cells if cell]
            line = " / ".join(cells)
        line = re.sub(r"^\s*[-*]\s+", "• ", line)
        output.append(clean_text(line))
    return output


def read_docx_text(path: Path) -> list[str]:
    doc = Document(str(path))
    lines = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    for table in doc.tables:
        for row in table.rows:
            values = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if values:
                lines.append(" | ".join(values))
    return lines


def read_pptx_text(path: Path) -> list[tuple[str, list[str]]]:
    if not path.exists():
        return []
    slides = []
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    with ZipFile(path) as archive:
        names = sorted(
            [name for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")],
            key=lambda name: int(re.search(r"slide(\d+)\.xml", name).group(1)),
        )
        for index, name in enumerate(names, start=1):
            root = ET.fromstring(archive.read(name))
            texts = [clean_text(node.text) for node in root.findall(".//a:t", ns) if node.text and clean_text(node.text)]
            slides.append((f"שקופית {index}", texts))
    return slides


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.footer_distance = Inches(0.35)
    set_page_direction(section)
    set_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = CLAY
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in [(1, 18), (2, 15), (3, 12.5)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INK if level == 1 else PURPLE if level == 2 else MAGENTA_DEEP
        style.paragraph_format.space_before = Pt(16 if level == 1 else 9)
        style.paragraph_format.space_after = Pt(5)

    for list_style_name in ["List Bullet", "List Number"]:
        style = styles[list_style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(10.5)
        style.font.color.rgb = CLAY
        style.paragraph_format.space_after = Pt(3)


def add_cover(doc: Document, data: dict) -> None:
    p = doc.add_paragraph()
    rtl(p)
    run = p.add_run("מסמך לאישור תוכן")
    set_run_font(run, size=10.5, bold=True, color=PURPLE)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    rtl(p)
    run = p.add_run("בארי אריזות")
    set_run_font(run, size=32, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    rtl(p)
    run = p.add_run("מסמך סקירת תוכן לאישור לקוח")
    set_run_font(run, size=18, bold=True, color=MAGENTA_DEEP)
    paragraph_border_bottom(p, color="00FFFF", size="22")
    p.paragraph_format.space_after = Pt(18)

    add_paragraph(
        doc,
        "מסמך זה מרכז את התוכן העברי הפעיל באתר, את המיקרו-קופי של הטפסים והמצבים, ואת עיקרי מסמכי האסטרטגיה והעבודה הרלוונטיים לאישור תוכן.",
        size=12,
    )
    add_callout(
        doc,
        "מטרת המסמך",
        "לאפשר ללקוח לעבור על כל התוכן במקום אחד: לאשר ניסוחים, לדייק מונחים, לבדוק נתונים עובדתיים ולסמן תיקונים לפני עלייה לאוויר.",
        fill=FILL_GOLD,
    )
    add_table(
        doc,
        ["שדה", "פירוט"],
        [
            ["שם חברה", data["company"]["legalNameHe"]],
            ["ח.פ.", data["company"]["registrationNumber"]],
            ["כתובת", f'{data["company"]["address"]["he"]["street"]}, {data["company"]["address"]["he"]["city"]}'],
            ["אימייל", data["company"]["email"]],
            ["תאריך הפקה", "8 ביוני 2026"],
            ["שפה", "עברית, RTL"],
        ],
    )
    add_heading(doc, "איך מומלץ לסקור את המסמך", 2)
    add_bullets(
        doc,
        [
            "לאשר קודם את המונחים המרכזיים, כי הם חוזרים בכל האתר.",
            "לעבור עמוד-עמוד ולסמן ניסוחים לתיקון, מחיקה או הרחבה.",
            "לבדוק במיוחד נתונים עובדתיים: שנים, תקנים, מספרי חברה, שטחים, שמות תפקידים ופרטי קשר.",
            "הפרקים המשפטיים מיועדים לסקירה עסקית בלבד ומומלץ להעביר אותם גם לעורך דין ישראלי לפני עלייה לאוויר.",
        ],
    )


def add_source_map(doc: Document) -> None:
    add_heading(doc, "מפת מקורות", 1, page_break=True)
    rows = [
        ["תוכן אתר פעיל", "app/content/*.ts, messages/he.json", "מופיע בפרקי העמודים והמיקרו-קופי"],
        ["תוכן בית מקור", "docs/content/homepage-content-he.md", "מופיע בנספח מקור התוכן לעמוד הבית"],
        ["מחקר SEO", "docs/research/beeri-google-trends-seo-research.md", "מופיע בנספח SEO"],
        ["אסטרטגיית תוכן", "docs/strategy/beeri-packaging-seo-content-strategy.docx", "מופיע בנספח אסטרטגיה"],
        ["תקציר בלוג SEO", "docs/briefs/beeri-seo-blog-brief.docx", "מופיע בנספח בלוג"],
        ["מצגת SEO", "docs/presentations/beeri-arizot-seo-findings.pptx", "מופיעה כטקסט שקופיות בנספח"],
        ["מסמכי תכנון/יישום", "docs/superpowers/specs + docs/superpowers/plans", "מופיעים כתקצירי אישור בעברית"],
    ]
    add_table(doc, ["סוג מקור", "קובץ", "אופן שימוש במסמך"], rows)


def add_terminology(doc: Document) -> None:
    add_heading(doc, "מילון מונחים לאישור", 1, page_break=True)
    add_paragraph(doc, "המונחים הבאים חוזרים באתר ובמסמכי התוכן. מומלץ לאשר אותם לפני עריכת ניסוחים נקודתיים.")
    add_table(doc, ["מונח", "הסבר לאישור"], TERMS)


def add_brand_and_chrome(doc: Document, data: dict) -> None:
    add_heading(doc, "מותג, מטא וניווט", 1, page_break=True)
    company = data["company"]
    metadata = data["metadata"]
    chrome = data["chrome"]

    add_heading(doc, "פרטי חברה", 2)
    add_table(
        doc,
        ["שדה", "תוכן"],
        [
            ["שם משפטי", company["legalNameHe"]],
            ["שם מותג", company["nameHe"]],
            ["קבוצה", company["groupHe"]],
            ["רישום חברה", str(company["foundingYear"])],
            ["ח.פ.", company["registrationNumber"]],
            ["כתובת", f'{company["address"]["he"]["street"]}, {company["address"]["he"]["city"]}, {company["address"]["he"]["country"]}'],
            ["אימייל", company["email"]],
            ["LinkedIn", company["linkedin"]],
        ],
    )

    add_heading(doc, "SEO בסיסי / Metadata", 2)
    add_label_value(doc, "שם אתר", metadata["siteName"])
    add_label_value(doc, "Title", metadata["title"])
    add_label_value(doc, "Description", metadata["description"])

    add_heading(doc, "ניווט ופוטר", 2)
    add_table(doc, ["פריט", "קישור"], [(item["he"], item["href"]) for item in data["navLinks"]])
    add_label_value(doc, "כפתור תפריט", chrome["menu"])
    add_label_value(doc, "כפתור סגירה", chrome["close"])
    add_label_value(doc, "כפתור שפה", chrome["lang"])
    add_label_value(doc, "כפתור יצירת קשר", chrome["contact"])
    add_label_value(doc, "כותרת פוטר", chrome["wordmark"])
    add_label_value(doc, "שורת פוטר", chrome["footerTagline"])
    add_label_value(doc, "מורשת", chrome["footerHeritage"])
    add_label_value(doc, "זכויות יוצרים", chrome["footerCopy"])


def add_home(doc: Document, data: dict) -> None:
    home = data["home"]
    copy = home["copy"]
    add_heading(doc, "עמוד הבית /he", 1, page_break=True)
    add_copy_block(
        doc,
        "הירו ופתיחה",
        [
            ("Eyebrow", copy["eyebrow"]),
            ("H1", join_title(copy["h1"])),
            ("CTA ראשי", copy["cta1"]),
            ("CTA משני", copy["cta2"]),
            ("הנחיית גלילה", copy["scroll"]),
            ("יצירת קשר", copy["contact"]),
        ],
    )
    add_copy_block(
        doc,
        "סיפור גלילה",
        [
            ("Eyebrow", copy["journeyEyebrow"]),
            ("כותרת", copy["journeyTitle"]),
            ("תיאור", copy["journeyDesc"]),
        ],
    )
    for panel in home["journeyPanels"]:
        add_copy_block(doc, panel["title"], [("תג", panel["tag"]), ("גוף", panel["body"]), ("קישור", panel["link"])])

    add_copy_block(
        doc,
        "יכולות וטכנולוגיה",
        [
            ("כותרת", copy["techTitle"]),
            ("גוף", copy["techBody"]),
            ("כרטיס 1", f'{copy["bento1Title"]} – {copy["bento1Body"]}'),
            ("כרטיס 2", f'{copy["bento2Title"]} – {copy["bento2Body"]}'),
            ("תג 1", copy["badge1"]),
            ("תג 2", copy["badge2"]),
        ],
    )
    add_table(doc, ["מספר", "יכולת", "טקסט"], [(item["n"], item["title"], item["body"]) for item in home["capabilities"]])

    add_copy_block(
        doc,
        "שאלות נפוצות",
        [("Eyebrow", copy["faqEyebrow"]), ("כותרת", copy["faqTitle"]), ("פתיחה", copy["faqBody"])],
    )
    for item in home["faqItems"]:
        add_copy_block(doc, f'שאלה {item["n"]}', [("שאלה", item["q"]), ("תשובה", item["a"])])
    add_label_value(doc, "CTA מסכם", join_title(copy["ctaTitle"]))


def add_about(doc: Document, data: dict) -> None:
    about = data["about"]
    add_heading(doc, "עמוד אודות /he/about", 1, page_break=True)
    add_copy_block(doc, "הירו", [("Eyebrow", about["eyebrow"]), ("כותרת", join_title(about["title"])), ("פתיחה", about["intro"])])
    add_copy_block(
        doc,
        "קבוצת דפוס בארי",
        [
            ("Eyebrow", about["heritageEyebrow"]),
            ("כותרת", about["heritageTitle"]),
            ("גוף", about["heritageBody"]),
            ("כיתוב תמונה", about["heritageImageCaption"]),
            ("קישור", f'{about["groupLinkLabel"]} – {about["groupLinkHref"]}'),
        ],
    )
    add_table(doc, ["שנה", "שם", "תיאור"], [(item["year"], item["name"], item["body"]) for item in about["heritageItems"]])
    add_copy_block(doc, "אבני דרך", [("Eyebrow", about["timelineEyebrow"]), ("כותרת", about["timelineTitle"])])
    add_table(doc, ["שנה", "כותרת", "תיאור"], [(item["year"], item["title"], item["body"]) for item in about["milestones"]])
    add_copy_block(doc, "מפעל ונתונים", [("Eyebrow", about["statsEyebrow"]), ("כותרת", about["statsTitle"])])
    add_table(doc, ["ערך", "מדד", "הערה"], [(item["value"], item["label"], item.get("sub", "")) for item in about["stats"]])
    add_copy_block(doc, "יכולות", [("Eyebrow", about["capsEyebrow"]), ("כותרת", about["capsTitle"]), ("גוף", about["capsBody"])])
    add_table(doc, ["שלב", "יכולת", "תיאור"], [(item["step"], item["title"], item["body"]) for item in about["capabilities"]])
    add_label_value(doc, "CTA השבחות", about["capsFinishingCta"])
    add_label_value(doc, "CTA קטלוג", about["capsCatalogCta"])
    add_copy_block(doc, "לקוחות ואיכות", [("לקוחות Eyebrow", about["partnersEyebrow"]), ("כותרת לקוחות", about["partnersTitle"]), ("תקנים Eyebrow", about["qualityEyebrow"]), ("כותרת תקנים", about["qualityTitle"])])
    add_table(doc, ["תקן", "כותרת", "תיאור"], [(item["code"], item["title"], item["body"]) for item in about["standards"]])
    add_copy_block(doc, "CTA", [("כותרת", about["ctaTitle"]), ("ראשי", about["ctaPrimary"]), ("משני", about["ctaSecondary"])])


def add_catalog(doc: Document, data: dict) -> None:
    catalog = data["catalog"]
    add_heading(doc, "קטלוג /he/catalog", 1, page_break=True)
    add_copy_block(doc, "הירו", [("Eyebrow", catalog["eyebrow"]), ("כותרת", join_title(catalog["title"])), ("פתיחה", catalog["intro"]), ("כרטיס מפרט", " / ".join(catalog["specCard"]))])
    for category in catalog["categories"]:
        add_heading(doc, f'{category["number"]} / {category["name"]}', 2)
        add_label_value(doc, "ספירת פריטים", category["count"])
        for item in category["items"]:
            fields = [
                ("סדרה", item.get("series")),
                ("שם", item["name"]),
                ("תיאור", item["description"]),
                ("תגיות", ", ".join(tag["label"] for tag in item.get("tags", []))),
                ("מפרטים", " | ".join(f'{spec["label"]}: {spec["value"]}' for spec in item.get("specs", []))),
                ("שכבת Hover", item.get("overlayLabel")),
                ("נתוני Hover", " | ".join(item.get("overlaySpecs", []))),
                ("CTA", item.get("cta")),
            ]
            add_copy_block(doc, item["name"], fields)


def add_finishing(doc: Document, data: dict) -> None:
    finishing = data["finishing"]
    add_heading(doc, "השבחות /he/finishing", 1, page_break=True)
    add_copy_block(doc, "הירו", [("שלב", finishing["step"]), ("כותרת", join_title(finishing["title"])), ("פתיחה", finishing["intro"])])
    for key, title in [("feature", "פויל"), ("deboss", "הבלטה ודיבוס"), ("texture", "מרקם ולכה")]:
        item = finishing[key]
        add_copy_block(doc, title, [("Eyebrow", item["eyebrow"]), ("כותרת", item["title"]), ("גוף", item["body"]), ("דוגמה", item.get("sample")), ("CTA", item.get("cta"))])
    add_heading(doc, finishing["metricsTitle"], 2)
    add_table(doc, ["מדד", "ערך"], [(item["label"], item["value"]) for item in finishing["metrics"]])
    add_label_value(doc, "ציטוט", finishing["quote"])
    add_label_value(doc, "ייחוס", finishing["quoteBy"])
    add_copy_block(doc, "CTA", [("כותרת", finishing["ctaTitle"]), ("ראשי", finishing["ctaPrimary"]), ("משני", finishing["ctaSecondary"]), ("כרטיס דוגמה", f'{finishing["sampleCard"]["value"]} – {finishing["sampleCard"]["label"]}'), ("כרטיס ISO", f'{finishing["isoCard"]["value"]} – {finishing["isoCard"]["label"]}')])


def add_careers(doc: Document, data: dict) -> None:
    careers = data["careers"]
    add_heading(doc, "קריירה /he/careers", 1, page_break=True)
    add_copy_block(doc, "הירו", [("Eyebrow", careers["eyebrow"]), ("כותרת", join_title(careers["title"])), ("פתיחה", careers["intro"]), ("Placeholder חיפוש", careers["searchPlaceholder"]), ("כפתור חיפוש", careers["searchButtonLabel"])])
    add_heading(doc, "כרטיסי תוכן", 2)
    for article in careers["articles"]:
        add_copy_block(doc, article["tag"], [("תאריך/מטא", article.get("meta")), ("כותרת", join_title(article["title"])), ("גוף", article.get("body")), ("CTA", article.get("cta"))])
    add_heading(doc, careers["rolesTitle"], 2)
    add_table(doc, ["קוד", "סטטוס", "תפקיד", "היקף", "מיקום", "מחלקה"], [(role["code"], role["status"], role["title"], role["scope"], role["location"], role["department"]) for role in careers["roles"]])
    add_label_value(doc, "כפתור הגשה", careers["apply"])
    add_label_value(doc, "אין משרות", careers["noRoles"])
    add_copy_block(doc, "עדכוני קריירה", [("כותרת", join_title(careers["newsletterTitle"])), ("גוף", careers["newsletterBody"]), ("Placeholder אימייל", careers["emailPlaceholder"]), ("CTA", careers["newsletterCta"]), ("הצלחה", careers["newsletterSuccess"]), ("שגיאה", careers["newsletterError"])])

    job = data["jobApplication"]
    add_heading(doc, "טופס הגשת מועמדות", 2)
    add_copy_block(doc, "צד מידע", [("Eyebrow", job["aside"]["eyebrow"]), ("Kicker", job["aside"]["kicker"]), ("Lead", job["aside"]["lead"]), ("תווית קוד", job["aside"]["codeLabel"]), ("כותרת ברירת מחדל", job["aside"]["fallbackTitle"]), ("פנייה כללית", job["aside"]["inquire"])])
    add_table(doc, ["יתרון", "גוף"], [(item["title"], item["body"]) for item in job["aside"]["perks"]])
    add_copy_block(doc, "שדות טופס", [("כותרת", job["form"]["heading"]), ("פתיח", job["form"]["intro"]), ("תווית משרה", job["form"]["roleLabel"]), ("שם", f'{job["form"]["name"]["label"]} / {job["form"]["name"]["placeholder"]}'), ("טלפון", f'{job["form"]["phone"]["label"]} / {job["form"]["phone"]["placeholder"]}'), ("דוא\"ל", f'{job["form"]["email"]["label"]} / {job["form"]["email"]["placeholder"]}'), ("הודעה", f'{job["form"]["message"]["label"]} / {job["form"]["message"]["placeholder"]}'), ("קורות חיים", f'{job["form"]["cv"]["label"]} / {job["form"]["cv"]["button"]} / {job["form"]["cv"]["empty"]}'), ("שליחה", job["form"]["submit"]), ("מצב שליחה", job["form"]["sending"])])
    add_table(doc, ["מצב", "טקסט"], list(job["errors"].items()) + [(f'הצלחה – {k}', v) for k, v in job["success"].items()])


def add_blog(doc: Document, data: dict) -> None:
    blog = data["blog"]
    index = blog["index"]
    add_heading(doc, "יומן / בלוג /he/blog", 1, page_break=True)
    add_copy_block(doc, "עמוד אינדקס", [("Eyebrow", index["eyebrow"]), ("כותרת", join_title(index["title"])), ("Lead", index["lead"]), ("גוף", index["body"]), ("הערת בקרוב", index["comingSoon"]), ("קריאה", index["readMore"]), ("חזרה", index["backToBlog"]), ("פורסם", index["publishedOn"]), ("404 פוסט – כותרת", index["notFoundTitle"]), ("404 פוסט – גוף", index["notFoundBody"])])
    add_heading(doc, "פוסטים", 2)
    for post in blog["posts"]:
        add_heading(doc, post["title"], 3)
        add_label_value(doc, "Slug", post["slug"])
        add_label_value(doc, "תאריך", post["date"])
        add_label_value(doc, "זמן קריאה", post["read"])
        add_label_value(doc, "קטגוריה", post["category"])
        add_label_value(doc, "תקציר", post["excerpt"])
        for paragraph in post.get("body", []):
            add_paragraph(doc, paragraph)
        if post.get("quote"):
            add_label_value(doc, "ציטוט", post["quote"]["text"])
            add_label_value(doc, "ייחוס", post["quote"]["cite"])
        for section in post.get("sections", []):
            add_copy_block(doc, section["heading"], [("גוף", section["body"])])


def add_contact_and_system(doc: Document, data: dict) -> None:
    contact = data["contact"]
    add_heading(doc, "יצירת קשר ומצבי מערכת", 1, page_break=True)
    add_copy_block(doc, "טופס יצירת קשר", [("סגירה", contact["closeLabel"]), ("Eyebrow", contact["eyebrow"]), ("כותרת", contact["title"]), ("תיאור נגישות", contact["description"])])
    form_rows = []
    for key in ["fullName", "phone", "email", "company", "details"]:
        item = contact["form"][key]
        form_rows.append((key, item["label"], item["placeholder"]))
    add_table(doc, ["שדה", "תווית", "Placeholder"], form_rows)
    add_table(doc, ["ערך", "תווית"], [(item["value"], item["label"]) for item in contact["form"]["reason"]["options"]])
    add_label_value(doc, "כפתור שליחה", contact["form"]["submit"])
    add_label_value(doc, "מצב שליחה", contact["form"]["sending"])
    add_label_value(doc, "הסכמה", contact["form"]["consent"])
    add_table(doc, ["מצב", "טקסט"], list(contact["errors"].items()) + [(f'הצלחה – {k}', v) for k, v in contact["success"].items()])

    system = data["system"]
    for title, key in [("עמוד 404", "notFound"), ("עמוד שגיאה", "error")]:
        item = system[key]
        add_copy_block(doc, title, [("Eyebrow", item["eyebrow"]), ("קוד", item.get("code")), ("כותרת", item["title"]), ("תיאור", item["description"]), ("CTA ראשי", item["primary"]), ("CTA משני", item["secondary"])])


def add_legal(doc: Document, data: dict) -> None:
    add_heading(doc, "מסמכים משפטיים", 1, page_break=True)
    for label, legal_doc in [("מדיניות פרטיות /he/privacy", data["legal"]["privacy"]), ("תנאי שימוש /he/terms", data["legal"]["terms"])]:
        add_heading(doc, label, 2)
        add_label_value(doc, "Eyebrow", legal_doc["eyebrow"])
        add_label_value(doc, "כותרת", legal_doc["title"])
        add_label_value(doc, "עדכון", legal_doc["updated"])
        for paragraph in legal_doc["intro"]:
            add_paragraph(doc, paragraph)
        for section in legal_doc["sections"]:
            add_heading(doc, section["heading"], 3)
            for paragraph in section.get("body", []):
                add_paragraph(doc, paragraph)
            add_bullets(doc, section.get("list", []))
        add_label_value(doc, legal_doc["contactHeading"], legal_doc["contactIntro"])


def add_placeholders(doc: Document, data: dict) -> None:
    add_heading(doc, "תוכן fallback / placeholder", 1, page_break=True)
    add_paragraph(doc, "התוכן הבא קיים בקוד כתוכן גיבוי למסכים שעדיין לא נמשכים ממערכת התוכן. חלקו עשוי לא להופיע כיום באתר החי אם העמוד כבר הוחלף בתוכן מלא.")
    for route, item in data["placeholder"].items():
        add_copy_block(doc, route, [("Eyebrow", item["eyebrow"]), ("כותרת", join_title(item["title"])), ("Lead", item["lead"]), ("גוף", item["body"]), ("Preview", " | ".join(item["preview"])), ("CTA ראשי", item["ctaPrimary"]), ("CTA משני", item["ctaSecondary"])])


def add_docs_appendix(doc: Document) -> None:
    add_heading(doc, "נספח מסמכי תוכן ואסטרטגיה", 1, page_break=True)
    add_heading(doc, "מסמך מקור: תוכן לעמוד הבית", 2)
    for line in read_markdown_text(ROOT / "docs" / "content" / "homepage-content-he.md"):
        if line.startswith("• "):
            add_bullets(doc, [line[2:]])
        else:
            add_paragraph(doc, line)

    add_heading(doc, "מחקר SEO", 2, page_break=True)
    for line in read_markdown_text(ROOT / "docs" / "research" / "beeri-google-trends-seo-research.md"):
        if line.startswith("http"):
            add_label_value(doc, "קישור מחקר", line)
        elif line.startswith("• "):
            add_bullets(doc, [line[2:]])
        else:
            add_paragraph(doc, line)

    for source in PYTHON_DOCX_SOURCES:
        add_heading(doc, source.relative_to(ROOT).as_posix(), 2, page_break=True)
        for line in read_docx_text(source):
            add_paragraph(doc, line)

    add_heading(doc, "מצגת SEO – טקסט שקופיות", 2, page_break=True)
    for slide_title, lines in read_pptx_text(PPTX_SOURCE):
        add_heading(doc, slide_title, 3)
        for line in lines:
            add_paragraph(doc, line)

    add_heading(doc, "תקצירי מסמכי תכנון ויישום", 2, page_break=True)
    for item in DOC_SUMMARIES:
        add_heading(doc, item["title"], 3)
        add_label_value(doc, "קובץ מקור", item["file"])
        add_bullets(doc, item["summary"])


def build_document() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = export_site_data()
    doc = Document()
    configure_document(doc)

    add_cover(doc, data)
    add_source_map(doc)
    add_terminology(doc)
    add_brand_and_chrome(doc, data)
    add_home(doc, data)
    add_about(doc, data)
    add_catalog(doc, data)
    add_finishing(doc, data)
    add_careers(doc, data)
    add_blog(doc, data)
    add_contact_and_system(doc, data)
    add_legal(doc, data)
    add_placeholders(doc, data)
    add_docs_appendix(doc)

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_document())
